
# %%
import requests
from bs4 import BeautifulSoup
from PIL import Image
import re
import codecs
import os
import pickle
from langchain.document_loaders import UnstructuredURLLoader
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.vectorstores import Chroma
from langchain.docstore.document import Document
from langchain.prompts import PromptTemplate
from langchain.indexes.vectorstore import VectorstoreIndexCreator
from urllib.request import urlretrieve
import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# 请替换为你的OpenAI API Key
os.environ["OPENAI_API_KEY"] = 'xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx' 
# 论文网页URL，目前只在校园网环境下，对Nature及其子刊有效  
url = "https://www.nature.com/articles/s41558-023-01642-3"

# 从网页中提取论文-文字内容
loaders = UnstructuredURLLoader(urls=[url])
data = loaders.load()
# Text Splitter
text_splitter = CharacterTextSplitter( separator = "\n",
                                      chunk_size=1000, 
                                      chunk_overlap=10)
texts = text_splitter.split_documents(data)
texts = [text.page_content for text in texts] #不这么把文字从doc中提取出来，Chroma.from_texts会报错
embeddings = OpenAIEmbeddings()
docsearch = Chroma.from_texts(texts, embeddings, metadatas=[{"source": str(i)} for i in range(len(texts))]).as_retriever()

#%%
# 设置请求头，模拟浏览器发送请求
headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3'}
response = requests.get(url, headers=headers)

# 解析网页内容，提取包含论文-插图的元素
soup = BeautifulSoup(response.content, 'html.parser')
elements = soup.find_all('div', {'class': 'c-article-section__figure js-c-reading-companion-figures-item'})
doc = docx.Document()

# 构建openai的QA模型
chain = load_qa_chain(OpenAI(temperature=0), chain_type="stuff")
query = "What does fig 1 show?"

def ask_question(query):
    docs = docsearch.get_relevant_documents(query)
    reply = chain({"input_documents": docs, "question": query}, return_only_outputs=True)
    return reply['output_text']

#创建引文信息
query_bio = "Create bibliography for this paper. Including the doi number if possibe."
answer_bio = ask_question(query)

#用于总结文章的问题
queries = ['What are questions that this paper solves?',
           'please summarise the Methods of this paper.',
           'What are the main findings of this paper?',
           'What are the implications of this paper?']
answers = [ ask_question(query) for query in queries ]




queries_fig = [f'What does fig {img_num} show?' for img_num in range(1, len(elements)+1)]
answers_fig = [ ask_question(query) for query in queries_fig ]
#%%
# 创建一个word容器
doc = docx.Document()

# 写入引文信息
p = doc.add_paragraph()
p.add_run(answer_bio.strip())

# 写入总结信息
for i, answer in enumerate(answers):
    p = doc.add_paragraph()
    p.add_run(answer.strip())
# 循环遍历每个元素，提取其中的图片,并和图片的总结写入word
img_num = 1
for i, element in enumerate(elements):
    # 提取图片
    img_url = 'https:'+element.find('picture').find('source')['srcset']
    img_response = requests.get(img_url, stream=True)
    img = Image.open(img_response.raw)
    # 保存图片
    if os.path.exists(f'img{img_num}.png'):
        os.remove(f'img{img_num}.png')
    img.save(f'img{img_num}.png')
    
    # 写入图片到word，并居中
    doc.add_picture(f'img{img_num}.png', width=Inches(6))   
    last_paragraph = doc.paragraphs[-1] 
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    

    # 写入图片的总结
    p = doc.add_paragraph()
    p.add_run(answers_fig[img_num-1].strip()) # img_num-1是因为answers_fig的索引从0开始，而elements的索引从1开始
    img_num += 1
# 保存文本到word
file_out = 'output_Nature.docx'
if os.path.exists(file_out):
    os.remove(file_out)
doc.save(file_out)